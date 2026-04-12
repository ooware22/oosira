# -*- coding: utf-8 -*-
"""
CV Analyzer - Parser de documents texte
Gere DOCX, DOC, ODT, RTF, TXT.
"""

import os
import time
import shutil
import logging
import tempfile
import subprocess
import sys
from pathlib import Path

logger = logging.getLogger("cv_analyzer.docx_parser")


class DocumentParser:
    """Extraction de contenu depuis les documents texte."""

    def __init__(self, libreoffice_timeout: int = 30):
        self.libreoffice_timeout = libreoffice_timeout
        self._libreoffice_path = self._find_libreoffice()

    def extract(self, file_path: Path) -> dict:
        """
        Extrait le contenu d'un document.

        Routing selon l'extension :
        - .docx -> python-docx
        - .doc / .odt -> conversion LibreOffice puis python-docx
        - .rtf -> striprtf
        - .txt -> lecture directe

        Returns:
            dict avec text, paragraphs, tables, styles, metadata, etc.
        """
        t0 = time.perf_counter()
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"Fichier introuvable : {file_path}")

        ext = file_path.suffix.lower()

        result = {
            "text": "",
            "paragraphs": [],
            "tables": [],
            "styles": {},
            "metadata": {},
            "processing_time_ms": 0,
            "conversion_used": False,
            "warnings": []
        }

        try:
            if ext == ".docx":
                result = self._parse_docx(file_path, result)
            elif ext in (".doc", ".odt"):
                result = self._parse_via_conversion(file_path, result)
            elif ext == ".rtf":
                result = self._parse_rtf(file_path, result)
            elif ext == ".txt":
                result = self._parse_txt(file_path, result)
            else:
                result["warnings"].append(f"Extension non supportee : {ext}")
        except Exception as e:
            logger.error(f"Erreur extraction {file_path.name} : {e}")
            result["warnings"].append(f"Erreur : {e}")

        # Assembler le texte complet si pas deja fait
        if not result["text"] and result["paragraphs"]:
            result["text"] = "\n".join(result["paragraphs"])

        result["processing_time_ms"] = int((time.perf_counter() - t0) * 1000)

        logger.info(
            f"Document extrait : {file_path.name}, {len(result['text'])} chars, "
            f"{len(result['paragraphs'])} paragraphes, "
            f"{len(result['tables'])} tableaux, "
            f"{result['processing_time_ms']} ms"
        )

        return result

    # -----------------------------------------------------------------
    # DOCX
    # -----------------------------------------------------------------

    def _parse_docx(self, file_path: Path, result: dict) -> dict:
        """Parse un fichier DOCX via python-docx."""
        try:
            from docx import Document
        except ImportError:
            result["warnings"].append("python-docx non installe")
            return result

        try:
            doc = Document(str(file_path))
        except Exception as e:
            result["warnings"].append(f"Fichier DOCX corrompu ou invalide : {e}")
            return result

        # Paragraphes avec styles
        style_counts = {}
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                result["paragraphs"].append(text)

                style_name = para.style.name if para.style else "Normal"
                if style_name not in style_counts:
                    style_counts[style_name] = 0
                style_counts[style_name] += 1

        result["styles"] = style_counts

        # Tableaux
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                result["tables"].append(table_data)

        # Texte complet (paragraphes + contenu des tableaux)
        all_texts = list(result["paragraphs"])
        for table in result["tables"]:
            for row in table:
                for cell in row:
                    if cell.strip():
                        all_texts.append(cell)
        result["text"] = "\n".join(all_texts)

        # Metadonnees
        try:
            props = doc.core_properties
            result["metadata"] = {
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
                "last_modified_by": props.last_modified_by or "",
                "category": props.category or "",
            }
        except Exception:
            pass

        return result

    # -----------------------------------------------------------------
    # DOC / ODT (via conversion LibreOffice)
    # -----------------------------------------------------------------

    def _parse_via_conversion(self, file_path: Path, result: dict) -> dict:
        """Convertir en DOCX via LibreOffice puis parser."""
        result["conversion_used"] = True

        if not self._libreoffice_path:
            result["warnings"].append(
                "LibreOffice non installe. Impossible de convertir les fichiers "
                f"{file_path.suffix}. Installer LibreOffice : https://www.libreoffice.org/"
            )
            # Tenter une lecture texte brute en fallback
            return self._parse_txt_fallback(file_path, result)

        # Conversion dans un dossier temporaire
        with tempfile.TemporaryDirectory() as temp_dir:
            try:
                cmd = [
                    self._libreoffice_path,
                    "--headless", "--convert-to", "docx",
                    "--outdir", temp_dir,
                    str(file_path)
                ]

                logger.info(f"Conversion LibreOffice : {file_path.name} -> DOCX")
                proc = subprocess.run(
                    cmd,
                    capture_output=True, text=True,
                    timeout=self.libreoffice_timeout,
                    creationflags=subprocess.CREATE_NO_WINDOW if sys.platform == "win32" else 0
                )

                if proc.returncode != 0:
                    result["warnings"].append(f"LibreOffice conversion echouee : {proc.stderr}")
                    return self._parse_txt_fallback(file_path, result)

                # Trouver le fichier converti
                converted_files = list(Path(temp_dir).glob("*.docx"))
                if not converted_files:
                    result["warnings"].append("Aucun fichier DOCX genere par LibreOffice")
                    return self._parse_txt_fallback(file_path, result)

                # Parser le DOCX converti
                result = self._parse_docx(converted_files[0], result)
                result["conversion_used"] = True
                return result

            except subprocess.TimeoutExpired:
                result["warnings"].append(
                    f"Timeout LibreOffice ({self.libreoffice_timeout}s). "
                    "Le fichier est peut-etre trop volumineux."
                )
                return self._parse_txt_fallback(file_path, result)

            except Exception as e:
                result["warnings"].append(f"Erreur conversion : {e}")
                return self._parse_txt_fallback(file_path, result)

    # -----------------------------------------------------------------
    # RTF
    # -----------------------------------------------------------------

    def _parse_rtf(self, file_path: Path, result: dict) -> dict:
        """Parse un fichier RTF."""
        try:
            from striprtf.striprtf import rtf_to_text
        except ImportError:
            result["warnings"].append("striprtf non installe, tentative de conversion LibreOffice")
            return self._parse_via_conversion(file_path, result)

        try:
            # Essayer plusieurs encodages
            for encoding in ["utf-8", "latin-1", "cp1252"]:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        rtf_content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                with open(file_path, "r", encoding="latin-1", errors="replace") as f:
                    rtf_content = f.read()

            text = rtf_to_text(rtf_content)
            result["text"] = text.strip()
            result["paragraphs"] = [p.strip() for p in text.split("\n") if p.strip()]

        except Exception as e:
            result["warnings"].append(f"Erreur parsing RTF : {e}")
            # Fallback LibreOffice
            return self._parse_via_conversion(file_path, result)

        return result

    # -----------------------------------------------------------------
    # TXT
    # -----------------------------------------------------------------

    def _parse_txt(self, file_path: Path, result: dict) -> dict:
        """Parse un fichier texte brut."""
        text = self._read_text_file(file_path)
        result["text"] = text.strip()
        result["paragraphs"] = [p.strip() for p in text.split("\n") if p.strip()]
        return result

    def _parse_txt_fallback(self, file_path: Path, result: dict) -> dict:
        """Tentative de lecture texte brute (fallback)."""
        try:
            text = self._read_text_file(file_path)
            if text.strip():
                result["text"] = text.strip()
                result["paragraphs"] = [p.strip() for p in text.split("\n") if p.strip()]
                result["warnings"].append("Fichier lu en mode texte brut (fallback)")
        except Exception as e:
            result["warnings"].append(f"Lecture texte brute echouee : {e}")
        return result

    # -----------------------------------------------------------------
    # Utilitaires
    # -----------------------------------------------------------------

    def _read_text_file(self, file_path: Path) -> str:
        """Lit un fichier texte en detectant l'encodage."""
        encodings = ["utf-8-sig", "utf-8", "latin-1", "cp1252", "iso-8859-1"]

        for enc in encodings:
            try:
                with open(file_path, "r", encoding=enc) as f:
                    content = f.read()
                # Verifier que le contenu est lisible
                if content and not all(c in '\x00\x01\x02\x03\x04\x05' for c in content[:100]):
                    logger.debug(f"Fichier lu avec encodage : {enc}")
                    return content
            except (UnicodeDecodeError, UnicodeError):
                continue

        # Dernier recours : latin-1 avec remplacement
        with open(file_path, "r", encoding="latin-1", errors="replace") as f:
            return f.read()

    def _find_libreoffice(self) -> str:
        """Cherche l'executable LibreOffice sur le systeme."""
        if sys.platform == "win32":
            candidates = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            ]
            # Chercher aussi dans les variables d'environnement
            lo_path = os.environ.get("LIBREOFFICE_PATH", "")
            if lo_path:
                candidates.insert(0, lo_path)

            for path in candidates:
                if os.path.isfile(path):
                    logger.info(f"LibreOffice trouve : {path}")
                    return path
        else:
            # Linux / Mac
            for cmd in ["libreoffice", "soffice"]:
                result = shutil.which(cmd)
                if result:
                    logger.info(f"LibreOffice trouve : {result}")
                    return result

        logger.info("LibreOffice non trouve sur le systeme")
        return ""
