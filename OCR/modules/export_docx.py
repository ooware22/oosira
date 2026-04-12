# -*- coding: utf-8 -*-
"""
CV Analyzer - Export DOCX
Genere des fichiers Word avec 3 templates : classic, modern, minimal.
"""

import logging
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

import config

logger = logging.getLogger("cv_analyzer.export_docx")


class DOCXExporter:
    """Genere des fichiers DOCX a partir du JSON CV normalise."""

    def export(self, cv_data: dict, template: str = "classic",
               output_path: Path = None) -> Path:
        """
        Genere un DOCX.

        Args:
            cv_data: JSON normalise du CV
            template: "classic" | "modern" | "minimal"
            output_path: chemin de sortie

        Returns:
            Path du fichier genere
        """
        if output_path is None:
            name = cv_data.get("personal_info", {}).get("full_name", "cv") or "cv"
            name = "".join(c for c in name if c.isalnum() or c in (" ", "-", "_")).strip()
            output_path = config.OUTPUTS_DIR / f"{name}_{template}.docx"

        output_path.parent.mkdir(parents=True, exist_ok=True)

        handlers = {
            "classic": self._build_classic,
            "modern": self._build_modern,
            "minimal": self._build_minimal,
        }

        handler = handlers.get(template, self._build_classic)
        doc = handler(cv_data)

        doc.save(str(output_path))
        logger.info(f"DOCX exporte : {output_path} (template={template})")
        return output_path

    # =================================================================
    # Template CLASSIC
    # =================================================================

    def _build_classic(self, cv: dict) -> Document:
        """Template classique : sobre, ATS-friendly."""
        doc = Document()
        info = cv.get("personal_info", {})

        # Marges 2cm
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        # Style par defaut
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # -- Nom --
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(info.get("full_name", ""))
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

        # -- Titre professionnel --
        title = cv.get("professional_title", "")
        if title:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title)
            run.italic = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # -- Contact --
        contact_parts = []
        if info.get("email"):
            contact_parts.append(info["email"])
        if info.get("phone"):
            contact_parts.append(info["phone"])
        if info.get("city"):
            loc = info["city"]
            if info.get("country"):
                loc += f", {info['country']}"
            contact_parts.append(loc)
        if info.get("linkedin"):
            contact_parts.append(info["linkedin"])

        if contact_parts:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(" | ".join(contact_parts))
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # -- Resume --
        summary = cv.get("summary", "")
        if summary:
            self._add_section_header(doc, "PROFIL")
            doc.add_paragraph(summary)

        # -- Experience --
        if cv.get("experience"):
            self._add_section_header(doc, "EXPERIENCE PROFESSIONNELLE")
            for exp in cv["experience"]:
                p = doc.add_paragraph()
                run = p.add_run(exp.get("job_title", ""))
                run.bold = True
                run.font.size = Pt(11)

                if exp.get("company"):
                    run = p.add_run(f"  |  {exp['company']}")
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

                dates = self._format_dates(exp)
                if dates:
                    run = p.add_run(f"  |  {dates}")
                    run.italic = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

                if exp.get("description"):
                    doc.add_paragraph(exp["description"])

                for ach in exp.get("achievements", []):
                    p = doc.add_paragraph(ach, style="List Bullet")
                    p.paragraph_format.space_after = Pt(2)

        # -- Formation --
        if cv.get("education"):
            self._add_section_header(doc, "FORMATION")
            for edu in cv["education"]:
                p = doc.add_paragraph()
                run = p.add_run(edu.get("degree", ""))
                run.bold = True

                if edu.get("institution"):
                    run = p.add_run(f"  |  {edu['institution']}")
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

                dates = self._format_dates(edu)
                if dates:
                    run = p.add_run(f"  |  {dates}")
                    run.italic = True
                    run.font.size = Pt(10)

        # -- Competences --
        skills = cv.get("skills", {})
        tech = skills.get("technical", [])
        soft = skills.get("soft", [])
        if tech or soft:
            self._add_section_header(doc, "COMPETENCES")
            if tech:
                doc.add_paragraph(f"Techniques : {', '.join(tech)}")
            if soft:
                doc.add_paragraph(f"Transversales : {', '.join(soft)}")

        # -- Langues --
        langs = skills.get("languages", [])
        if langs:
            self._add_section_header(doc, "LANGUES")
            for lang in langs:
                txt = lang.get("language", "")
                if lang.get("level"):
                    txt += f" - {lang['level']}"
                doc.add_paragraph(txt)

        # -- Centres d'interet --
        interests = cv.get("interests", [])
        if interests:
            self._add_section_header(doc, "CENTRES D'INTERET")
            doc.add_paragraph(", ".join(interests))

        # Pied de page
        self._add_page_number(doc)

        return doc

    # =================================================================
    # Template MODERN
    # =================================================================

    def _build_modern(self, cv: dict) -> Document:
        """Template moderne : sidebar coloree + corps."""
        doc = Document()
        info = cv.get("personal_info", {})

        for section in doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1.5)

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)

        # Table 2 colonnes : sidebar | corps
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # Largeurs colonnes
        sidebar_cell = table.cell(0, 0)
        body_cell = table.cell(0, 1)

        sidebar_cell.width = Cm(6)
        body_cell.width = Cm(12)

        # Fond sidebar
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1a365d" w:val="clear"/>')
        sidebar_cell._tc.get_or_add_tcPr().append(shading)

        # === SIDEBAR ===
        # Nom
        p = sidebar_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(info.get("full_name", ""))
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Titre
        title = cv.get("professional_title", "")
        if title:
            p = sidebar_cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xBB, 0xDE, 0xFB)

        # Contact
        self._sidebar_section(sidebar_cell, "CONTACT")
        if info.get("email"):
            self._sidebar_item(sidebar_cell, f"Email: {info['email']}")
        if info.get("phone"):
            self._sidebar_item(sidebar_cell, f"Tel: {info['phone']}")
        if info.get("city"):
            loc = info["city"]
            if info.get("country"):
                loc += f", {info['country']}"
            self._sidebar_item(sidebar_cell, loc)
        if info.get("linkedin"):
            self._sidebar_item(sidebar_cell, info["linkedin"])

        # Competences sidebar
        skills = cv.get("skills", {})
        tech = skills.get("technical", [])
        if tech:
            self._sidebar_section(sidebar_cell, "COMPETENCES")
            for skill in tech[:15]:
                self._sidebar_item(sidebar_cell, skill)

        # Langues sidebar
        langs = skills.get("languages", [])
        if langs:
            self._sidebar_section(sidebar_cell, "LANGUES")
            for lang in langs:
                txt = lang.get("language", "")
                if lang.get("level"):
                    txt += f" ({lang['level']})"
                self._sidebar_item(sidebar_cell, txt)

        # Interets sidebar
        interests = cv.get("interests", [])
        if interests:
            self._sidebar_section(sidebar_cell, "INTERETS")
            for interest in interests[:8]:
                self._sidebar_item(sidebar_cell, interest)

        # === CORPS ===
        # Resume
        summary = cv.get("summary", "")
        if summary:
            self._body_section(body_cell, "PROFIL", is_first=True)
            p = body_cell.add_paragraph(summary)
            p.paragraph_format.space_after = Pt(8)

        # Experience
        if cv.get("experience"):
            self._body_section(body_cell, "EXPERIENCE PROFESSIONNELLE")
            for exp in cv["experience"]:
                p = body_cell.add_paragraph()
                run = p.add_run(exp.get("job_title", ""))
                run.bold = True
                run.font.size = Pt(11)

                dates = self._format_dates(exp)
                if dates:
                    run = p.add_run(f"   {dates}")
                    run.italic = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

                if exp.get("company"):
                    p = body_cell.add_paragraph()
                    run = p.add_run(exp["company"])
                    run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
                    p.paragraph_format.space_after = Pt(2)

                if exp.get("description"):
                    p = body_cell.add_paragraph(exp["description"])
                    p.paragraph_format.space_after = Pt(4)

                for ach in exp.get("achievements", []):
                    p = body_cell.add_paragraph(f"- {ach}")
                    p.paragraph_format.space_after = Pt(1)
                    p.paragraph_format.left_indent = Cm(0.5)

        # Formation
        if cv.get("education"):
            self._body_section(body_cell, "FORMATION")
            for edu in cv["education"]:
                p = body_cell.add_paragraph()
                run = p.add_run(edu.get("degree", ""))
                run.bold = True

                dates = self._format_dates(edu)
                if dates:
                    run = p.add_run(f"   {dates}")
                    run.italic = True
                    run.font.size = Pt(9)

                if edu.get("institution"):
                    p = body_cell.add_paragraph()
                    run = p.add_run(edu["institution"])
                    run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

        # Supprimer les bordures du tableau
        self._remove_table_borders(table)
        self._add_page_number(doc)

        return doc

    # =================================================================
    # Template MINIMAL
    # =================================================================

    def _build_minimal(self, cv: dict) -> Document:
        """Template minimaliste : epure, typographie soignee."""
        doc = Document()
        info = cv.get("personal_info", {})

        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)
        style.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        # Nom
        p = doc.add_paragraph()
        run = p.add_run(info.get("full_name", ""))
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        run.font.name = "Georgia"
        p.paragraph_format.space_after = Pt(2)

        # Titre
        title = cv.get("professional_title", "")
        if title:
            p = doc.add_paragraph()
            run = p.add_run(title)
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x63, 0x66, 0xf1)
            p.paragraph_format.space_after = Pt(6)

        # Contact en ligne
        contact_parts = []
        for key in ["email", "phone", "linkedin"]:
            if info.get(key):
                contact_parts.append(info[key])
        if info.get("city"):
            contact_parts.append(info["city"])

        if contact_parts:
            p = doc.add_paragraph()
            run = p.add_run(" | ".join(contact_parts))
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            p.paragraph_format.space_after = Pt(16)

        # Separateur fin
        self._add_thin_line(doc)

        # Resume
        summary = cv.get("summary", "")
        if summary:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run(summary)
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            p.paragraph_format.space_after = Pt(12)

        # Experience
        if cv.get("experience"):
            self._minimal_section(doc, "Experience")
            for exp in cv["experience"]:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)

                run = p.add_run(exp.get("job_title", ""))
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

                dates = self._format_dates(exp)
                if exp.get("company") or dates:
                    p2 = doc.add_paragraph()
                    parts = []
                    if exp.get("company"):
                        parts.append(exp["company"])
                    if dates:
                        parts.append(dates)
                    run = p2.add_run(" | ".join(parts))
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                    p2.paragraph_format.space_after = Pt(3)

                if exp.get("description"):
                    p3 = doc.add_paragraph(exp["description"])
                    p3.paragraph_format.space_after = Pt(2)

        # Formation
        if cv.get("education"):
            self._minimal_section(doc, "Formation")
            for edu in cv["education"]:
                p = doc.add_paragraph()
                run = p.add_run(edu.get("degree", ""))
                run.bold = True
                run.font.size = Pt(11)

                if edu.get("institution"):
                    dates = self._format_dates(edu)
                    parts = [edu["institution"]]
                    if dates:
                        parts.append(dates)
                    p2 = doc.add_paragraph()
                    run = p2.add_run(" | ".join(parts))
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # Competences
        skills = cv.get("skills", {})
        all_skills = skills.get("technical", []) + skills.get("soft", [])
        if all_skills:
            self._minimal_section(doc, "Competences")
            doc.add_paragraph(", ".join(all_skills))

        # Langues
        langs = skills.get("languages", [])
        if langs:
            self._minimal_section(doc, "Langues")
            parts = []
            for lang in langs:
                txt = lang.get("language", "")
                if lang.get("level"):
                    txt += f" ({lang['level']})"
                parts.append(txt)
            doc.add_paragraph(" | ".join(parts))

        self._add_page_number(doc)
        return doc

    # =================================================================
    # Utilitaires
    # =================================================================

    def _add_section_header(self, doc, title):
        """Ajoute un titre de section avec ligne horizontale (classic)."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        # Bordure inferieure
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="888888"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

    def _minimal_section(self, doc, title):
        """Titre de section minimaliste."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(title.upper())
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x63, 0x66, 0xf1)
        run.font.name = "Calibri"
        run.bold = False
        # Lettre spacing via XML
        rPr = run._r.get_or_add_rPr()
        spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:val="60"/>')
        rPr.append(spacing)

    def _sidebar_section(self, cell, title):
        """Titre de section dans la sidebar."""
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xBB, 0xDE, 0xFB)

    def _sidebar_item(self, cell, text):
        """Element dans la sidebar."""
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xE0, 0xE0, 0xE0)

    def _body_section(self, cell, title, is_first=False):
        """Titre de section dans le corps."""
        if is_first:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    def _add_thin_line(self, doc):
        """Ligne horizontale fine."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="2" w:space="1" w:color="DDDDDD"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

    def _remove_table_borders(self, table):
        """Supprime toutes les bordures d'un tableau."""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)

    def _add_page_number(self, doc):
        """Ajoute le numero de page en pied de page."""
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            p = footer.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            fld = parse_xml(
                f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
            )
            run._r.append(fld)
            run2 = p.add_run(" PAGE ")
            run2._r.append(parse_xml(
                f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
            ))
            run3 = p.add_run()
            run3._r.append(parse_xml(
                f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
            ))

    def _format_dates(self, entry: dict) -> str:
        """Formate les dates d'une entree."""
        start = entry.get("start_date", "")
        end = entry.get("end_date", "")
        if start and end:
            return f"{start} - {end}"
        if start:
            return f"Depuis {start}"
        if end:
            return end
        return ""
